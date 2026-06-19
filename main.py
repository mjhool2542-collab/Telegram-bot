import os
import logging
import tempfile
import telebot
from groq import Groq
import google.generativeai as genai
import cohere
from docx import Document
import fitz

# مهم: نستخدم.get عشان ما يفشل البناء
BOT_TOKEN = os.environ.get('BOT_TOKEN', 'dummy')
GEMINI_KEY = os.environ.get('GEMINI_KEY', 'dummy')
COHERE_KEY = os.environ.get('COHERE_KEY', 'dummy')
GROQ_KEY = os.environ.get('GROQ_KEY', 'dummy')

# لو أي مفتاح وهمي، نوقف التشغيل ونطلع خطأ واضح
if 'dummy' in [BOT_TOKEN, GEMINI_KEY, COHERE_KEY, GROQ_KEY]:
    raise ValueError("ضيف الـ4 مفاتيح: BOT_TOKEN, GEMINI_KEY, COHERE_KEY, GROQ_KEY في Variables")

logging.basicConfig(level=logging.INFO)
bot = telebot.TeleBot(BOT_TOKEN)
groq_client = Groq(api_key=GROQ_KEY)
genai.configure(api_key=GEMINI_KEY)
co = cohere.Client(COHERE_KEY)

def translate_text(text, target_lang):
    prompt = f"Translate to {target_lang}. Keep formatting, tables, numbers. Only translate text:\n\n{text}"
    try:
        chat = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.1-70b-versatile",
            temperature=0.1
        )
        return chat.choices[0].message.content
    except Exception as e:
        logging.error(f"Groq failed: {e}")
        model = genai.GenerativeModel('gemini-1.5-pro')
        return model.generate_content(prompt).text

@bot.message_handler(commands=['start'])
def start(message):
    bot.reply_to(message, "أهلاً! أرسل لي ملف PDF أو Word وأترجمه لك عربي ↔ إنجليزي ⚡")

@bot.message_handler(content_types=['document'])
def handle_docs(message):
    try:
        file_info = bot.get_file(message.document.file_id)
        downloaded = bot.download_file(file_info.file_path)
        msg = bot.reply_to(message, "استلمت الملف... جاري استخراج النص 🚀")

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(downloaded)
            path = tmp.name

        # استخراج النص
        if message.document.mime_type == 'application/pdf':
            doc = fitz.open(path)
            text = "".join([page.get_text() for page in doc])
        else:
            doc = Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])

        if not text.strip():
            bot.edit_message_text("الملف فاضي أو ما قدرت أقرأه", msg.chat.id, msg.message_id)
            return

        # تحديد اللغة الهدف
        target = "Arabic" if any("\u0600" <= c <= "\u06FF" for c in text[:100]) else "English"
        bot.edit_message_text(f"جاري الترجمة إلى {target}...", msg.chat.id, msg.message_id)

        translated = translate_text(text, target)

        # حفظ ملف جديد
        new_doc = Document()
        new_doc.add_paragraph(translated)
        out_path = path.replace('.docx', '_translated.docx')
        new_doc.save(out_path)

        with open(out_path, 'rb') as f:
            bot.send_document(message.chat.id, f, caption=f"تمت الترجمة إلى {target} ✅")

        os.remove(path)
        os.remove(out_path)

    except Exception as e:
        bot.reply_to(message, f"خطأ: {e}")

print("Translation Bot is running...")
bot.polling(none_stop=True)
